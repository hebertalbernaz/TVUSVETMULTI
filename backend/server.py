from fastapi import FastAPI, APIRouter, HTTPException, UploadFile, File
from fastapi.responses import FileResponse
from dotenv import load_dotenv
from starlette.middleware.cors import CORSMiddleware
from motor.motor_asyncio import AsyncIOMotorClient
import os
import logging
from pathlib import Path
from pydantic import BaseModel, Field, ConfigDict
from typing import List, Optional, Dict, Any
import uuid
from datetime import datetime, timezone
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import base64
import io

ROOT_DIR = Path(__file__).parent
load_dotenv(ROOT_DIR / '.env')

# MongoDB connection
mongo_url = os.environ['MONGO_URL']
client = AsyncIOMotorClient(mongo_url)
db = client[os.environ['DB_NAME']]

# Create the main app without a prefix
app = FastAPI()

# Create a router with the /api prefix
api_router = APIRouter(prefix="/api")

# Ensure upload directories exist
UPLOAD_DIR = ROOT_DIR / "uploads"
UPLOAD_DIR.mkdir(exist_ok=True)
IMAGES_DIR = UPLOAD_DIR / "images"
IMAGES_DIR.mkdir(exist_ok=True)
REPORTS_DIR = UPLOAD_DIR / "reports"
REPORTS_DIR.mkdir(exist_ok=True)

# Models
class Patient(BaseModel):
    model_config = ConfigDict(extra="ignore")
    
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    name: str
    species: str  # "dog" or "cat"
    breed: str
    weight: float  # in kg
    size: str  # "small", "medium", "large"
    sex: str  # "male" or "female"
    is_neutered: bool = False
    owner_name: Optional[str] = None
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

class PatientCreate(BaseModel):
    name: str
    species: str
    breed: str
    weight: float
    size: str
    sex: str
    is_neutered: bool = False
    owner_name: Optional[str] = None

class OrganMeasurement(BaseModel):
    value: float
    unit: str  # "cm" or "mm"
    is_abnormal: bool = False

class OrganData(BaseModel):
    organ_name: str
    measurements: Dict[str, OrganMeasurement] = {}
    selected_findings: List[str] = []
    custom_notes: str = ""
    report_text: str = ""

class ExamImage(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    filename: str
    organ: Optional[str] = None
    path: str

class Exam(BaseModel):
    model_config = ConfigDict(extra="ignore")
    
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    patient_id: str
    exam_date: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))
    organs_data: List[OrganData] = []
    images: List[ExamImage] = []
    final_report: str = ""
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

class ExamCreate(BaseModel):
    patient_id: str
    exam_date: Optional[datetime] = None

class ExamUpdate(BaseModel):
    organs_data: Optional[List[OrganData]] = None
    final_report: Optional[str] = None

class TemplateText(BaseModel):
    model_config = ConfigDict(extra="ignore")
    
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    organ: str
    category: str  # "normal", "finding", "conclusion"
    text: str
    order: int = 0

class TemplateTextCreate(BaseModel):
    organ: str
    category: str
    text: str
    order: int = 0

class ReferenceValue(BaseModel):
    model_config = ConfigDict(extra="ignore")
    
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    organ: str
    measurement_type: str
    species: str
    size: str
    min_value: float
    max_value: float
    unit: str

class ReferenceValueCreate(BaseModel):
    organ: str
    measurement_type: str
    species: str
    size: str
    min_value: float
    max_value: float
    unit: str

class Settings(BaseModel):
    model_config = ConfigDict(extra="ignore")
    
    id: str = "global_settings"
    letterhead_path: Optional[str] = None
    clinic_name: Optional[str] = None
    clinic_address: Optional[str] = None
    veterinarian_name: Optional[str] = None
    crmv: Optional[str] = None

# Helper functions
def prepare_for_mongo(data: dict) -> dict:
    """Convert datetime objects to ISO strings for MongoDB storage"""
    for key, value in data.items():
        if isinstance(value, datetime):
            data[key] = value.isoformat()
        elif isinstance(value, list):
            for i, item in enumerate(value):
                if isinstance(item, dict):
                    value[i] = prepare_for_mongo(item)
    return data

def parse_from_mongo(item: dict) -> dict:
    """Convert ISO string timestamps back to datetime objects"""
    if item is None:
        return None
    for key, value in item.items():
        if key in ['created_at', 'exam_date'] and isinstance(value, str):
            item[key] = datetime.fromisoformat(value)
    return item

# Patient endpoints
@api_router.post("/patients", response_model=Patient)
async def create_patient(patient_data: PatientCreate):
    patient = Patient(**patient_data.model_dump())
    doc = prepare_for_mongo(patient.model_dump())
    await db.patients.insert_one(doc)
    return patient

@api_router.get("/patients", response_model=List[Patient])
async def get_patients():
    patients = await db.patients.find({}, {"_id": 0}).to_list(1000)
    return [Patient(**parse_from_mongo(p)) for p in patients]

@api_router.get("/patients/{patient_id}", response_model=Patient)
async def get_patient(patient_id: str):
    patient = await db.patients.find_one({"id": patient_id}, {"_id": 0})
    if not patient:
        raise HTTPException(status_code=404, detail="Patient not found")
    return Patient(**parse_from_mongo(patient))

@api_router.put("/patients/{patient_id}", response_model=Patient)
async def update_patient(patient_id: str, patient_data: PatientCreate):
    patient = Patient(id=patient_id, **patient_data.model_dump())
    doc = prepare_for_mongo(patient.model_dump())
    result = await db.patients.update_one({"id": patient_id}, {"$set": doc})
    if result.matched_count == 0:
        raise HTTPException(status_code=404, detail="Patient not found")
    return patient

@api_router.delete("/patients/{patient_id}")
async def delete_patient(patient_id: str):
    result = await db.patients.delete_one({"id": patient_id})
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Patient not found")
    return {"message": "Patient deleted successfully"}

# Exam endpoints
@api_router.post("/exams", response_model=Exam)
async def create_exam(exam_data: ExamCreate):
    exam_dict = exam_data.model_dump()
    if exam_dict.get('exam_date') is None:
        exam_dict['exam_date'] = datetime.now(timezone.utc)
    exam = Exam(**exam_dict)
    doc = prepare_for_mongo(exam.model_dump())
    await db.exams.insert_one(doc)
    return exam

@api_router.get("/exams", response_model=List[Exam])
async def get_exams(patient_id: Optional[str] = None):
    query = {"patient_id": patient_id} if patient_id else {}
    exams = await db.exams.find(query, {"_id": 0}).sort("exam_date", -1).to_list(1000)
    return [Exam(**parse_from_mongo(e)) for e in exams]

@api_router.get("/exams/{exam_id}", response_model=Exam)
async def get_exam(exam_id: str):
    exam = await db.exams.find_one({"id": exam_id}, {"_id": 0})
    if not exam:
        raise HTTPException(status_code=404, detail="Exam not found")
    return Exam(**parse_from_mongo(exam))

@api_router.put("/exams/{exam_id}", response_model=Exam)
async def update_exam(exam_id: str, exam_data: ExamUpdate):
    update_dict = {k: v for k, v in exam_data.model_dump().items() if v is not None}
    if update_dict:
        update_dict = prepare_for_mongo(update_dict)
        result = await db.exams.update_one({"id": exam_id}, {"$set": update_dict})
        if result.matched_count == 0:
            raise HTTPException(status_code=404, detail="Exam not found")
    
    exam = await db.exams.find_one({"id": exam_id}, {"_id": 0})
    return Exam(**parse_from_mongo(exam))

@api_router.delete("/exams/{exam_id}")
async def delete_exam(exam_id: str):
    result = await db.exams.delete_one({"id": exam_id})
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Exam not found")
    return {"message": "Exam deleted successfully"}

# Template text endpoints
@api_router.post("/templates", response_model=TemplateText)
async def create_template(template_data: TemplateTextCreate):
    template = TemplateText(**template_data.model_dump())
    doc = template.model_dump()
    await db.templates.insert_one(doc)
    return template

@api_router.get("/templates", response_model=List[TemplateText])
async def get_templates(organ: Optional[str] = None):
    query = {"organ": organ} if organ else {}
    templates = await db.templates.find(query, {"_id": 0}).sort("order", 1).to_list(1000)
    return [TemplateText(**t) for t in templates]

@api_router.put("/templates/{template_id}", response_model=TemplateText)
async def update_template(template_id: str, template_data: TemplateTextCreate):
    template = TemplateText(id=template_id, **template_data.model_dump())
    doc = template.model_dump()
    result = await db.templates.update_one({"id": template_id}, {"$set": doc})
    if result.matched_count == 0:
        raise HTTPException(status_code=404, detail="Template not found")
    return template

@api_router.delete("/templates/{template_id}")
async def delete_template(template_id: str):
    result = await db.templates.delete_one({"id": template_id})
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Template not found")
    return {"message": "Template deleted successfully"}

# Reference values endpoints
@api_router.post("/reference-values", response_model=ReferenceValue)
async def create_reference_value(ref_data: ReferenceValueCreate):
    ref_value = ReferenceValue(**ref_data.model_dump())
    doc = ref_value.model_dump()
    await db.reference_values.insert_one(doc)
    return ref_value

@api_router.get("/reference-values", response_model=List[ReferenceValue])
async def get_reference_values(
    organ: Optional[str] = None,
    species: Optional[str] = None,
    size: Optional[str] = None
):
    query = {}
    if organ:
        query["organ"] = organ
    if species:
        query["species"] = species
    if size:
        query["size"] = size
    
    ref_values = await db.reference_values.find(query, {"_id": 0}).to_list(1000)
    return [ReferenceValue(**r) for r in ref_values]

@api_router.put("/reference-values/{ref_id}", response_model=ReferenceValue)
async def update_reference_value(ref_id: str, ref_data: ReferenceValueCreate):
    ref_value = ReferenceValue(id=ref_id, **ref_data.model_dump())
    doc = ref_value.model_dump()
    result = await db.reference_values.update_one({"id": ref_id}, {"$set": doc})
    if result.matched_count == 0:
        raise HTTPException(status_code=404, detail="Reference value not found")
    return ref_value

@api_router.delete("/reference-values/{ref_id}")
async def delete_reference_value(ref_id: str):
    result = await db.reference_values.delete_one({"id": ref_id})
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Reference value not found")
    return {"message": "Reference value deleted successfully"}

# Settings endpoints
@api_router.get("/settings", response_model=Settings)
async def get_settings():
    settings = await db.settings.find_one({"id": "global_settings"}, {"_id": 0})
    if not settings:
        # Create default settings
        default_settings = Settings()
        await db.settings.insert_one(default_settings.model_dump())
        return default_settings
    return Settings(**settings)

@api_router.put("/settings", response_model=Settings)
async def update_settings(settings_data: Settings):
    doc = settings_data.model_dump()
    await db.settings.update_one(
        {"id": "global_settings"},
        {"$set": doc},
        upsert=True
    )
    return settings_data

# Image upload endpoint
@api_router.post("/exams/{exam_id}/images")
async def upload_exam_image(exam_id: str, file: UploadFile = File(...), organ: Optional[str] = None):
    # Verify exam exists
    exam = await db.exams.find_one({"id": exam_id})
    if not exam:
        raise HTTPException(status_code=404, detail="Exam not found")
    
    # Save file
    file_ext = Path(file.filename).suffix
    image_id = str(uuid.uuid4())
    filename = f"{image_id}{file_ext}"
    filepath = IMAGES_DIR / filename
    
    content = await file.read()
    with open(filepath, "wb") as f:
        f.write(content)
    
    # Create image record
    image = ExamImage(
        id=image_id,
        filename=filename,
        organ=organ,
        path=str(filepath)
    )
    
    # Update exam
    await db.exams.update_one(
        {"id": exam_id},
        {"$push": {"images": image.model_dump()}}
    )
    
    return image

@api_router.get("/images/{image_id}")
async def get_image(image_id: str):
    # Find the image
    exam = await db.exams.find_one({"images.id": image_id})
    if not exam:
        raise HTTPException(status_code=404, detail="Image not found")
    
    image = next((img for img in exam.get("images", []) if img["id"] == image_id), None)
    if not image:
        raise HTTPException(status_code=404, detail="Image not found")
    
    filepath = Path(image["path"])
    if not filepath.exists():
        raise HTTPException(status_code=404, detail="Image file not found")
    
    return FileResponse(filepath)

@api_router.delete("/exams/{exam_id}/images/{image_id}")
async def delete_exam_image(exam_id: str, image_id: str):
    exam = await db.exams.find_one({"id": exam_id})
    if not exam:
        raise HTTPException(status_code=404, detail="Exam not found")
    
    # Find and delete image file
    image = next((img for img in exam.get("images", []) if img["id"] == image_id), None)
    if image:
        filepath = Path(image["path"])
        if filepath.exists():
            filepath.unlink()
    
    # Remove from exam
    await db.exams.update_one(
        {"id": exam_id},
        {"$pull": {"images": {"id": image_id}}}
    )
    
    return {"message": "Image deleted successfully"}

# Export to DOCX endpoint
@api_router.get("/exams/{exam_id}/export")
async def export_exam_to_docx(exam_id: str):
    # Get exam and patient data
    exam = await db.exams.find_one({"id": exam_id}, {"_id": 0})
    if not exam:
        raise HTTPException(status_code=404, detail="Exam not found")
    
    patient = await db.patients.find_one({"id": exam["patient_id"]}, {"_id": 0})
    if not patient:
        raise HTTPException(status_code=404, detail="Patient not found")
    
    settings = await db.settings.find_one({"id": "global_settings"}, {"_id": 0})
    
    # Create document
    doc = Document()
    
    # Add letterhead if configured
    if settings and settings.get("clinic_name"):
        heading = doc.add_heading(settings["clinic_name"], level=1)
        heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        if settings.get("clinic_address"):
            addr = doc.add_paragraph(settings["clinic_address"])
            addr.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        if settings.get("veterinarian_name") or settings.get("crmv"):
            vet_info = f"{settings.get('veterinarian_name', '')} - CRMV: {settings.get('crmv', '')}"
            vet_para = doc.add_paragraph(vet_info)
            vet_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        doc.add_paragraph()  # Spacer
    
    # Add title
    title = doc.add_heading('LAUDO DE ULTRASSONOGRAFIA ABDOMINAL', level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph()
    
    # Patient information
    doc.add_heading('Dados do Paciente', level=2)
    doc.add_paragraph(f"Nome: {patient['name']}")
    doc.add_paragraph(f"Espécie: {'Canino' if patient['species'] == 'dog' else 'Felino'}")
    doc.add_paragraph(f"Raça: {patient['breed']}")
    doc.add_paragraph(f"Peso: {patient['weight']} kg")
    doc.add_paragraph(f"Porte: {patient['size'].capitalize()}")
    doc.add_paragraph(f"Sexo: {'Macho' if patient['sex'] == 'male' else 'Fêmea'}")
    if patient.get('is_neutered'):
        doc.add_paragraph("Paciente Castrado")
    if patient.get('owner_name'):
        doc.add_paragraph(f"Tutor: {patient['owner_name']}")
    
    exam_date = exam.get('exam_date')
    if isinstance(exam_date, str):
        exam_date = datetime.fromisoformat(exam_date)
    doc.add_paragraph(f"Data do Exame: {exam_date.strftime('%d/%m/%Y')}")
    doc.add_paragraph()
    
    # Organ findings
    doc.add_heading('Achados Ultrassonográficos', level=2)
    
    organs_data = exam.get('organs_data', [])
    for organ_data in organs_data:
        if organ_data.get('report_text'):
            doc.add_heading(organ_data['organ_name'], level=3)
            doc.add_paragraph(organ_data['report_text'])
            doc.add_paragraph()
    
    # Save document
    output_path = REPORTS_DIR / f"laudo_{exam_id}.docx"
    doc.save(str(output_path))
    
    return FileResponse(
        path=output_path,
        filename=f"laudo_{patient['name']}_{exam_date.strftime('%Y%m%d')}.docx",
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

# Initialize default data
@api_router.post("/initialize-defaults")
async def initialize_defaults():
    """Initialize default template texts and reference values"""
    
    # Check if already initialized
    existing_templates = await db.templates.count_documents({})
    if existing_templates > 0:
        return {"message": "Defaults already initialized"}
    
    # Default organs
    organs = [
        "Estômago", "Fígado", "Baço", "Rim Esquerdo", "Rim Direito",
        "Vesícula Urinária", "Adrenal Esquerda", "Adrenal Direita",
        "Duodeno", "Jejuno", "Cólon", "Ceco", "Íleo", "Linfonodos"
    ]
    
    # Default template texts for each organ
    templates = []
    for idx, organ in enumerate(organs):
        templates.extend([
            TemplateText(
                organ=organ,
                category="normal",
                text=f"{organ} com dimensões, contornos, ecogenicidade e ecotextura preservados.",
                order=idx * 10
            ),
            TemplateText(
                organ=organ,
                category="finding",
                text=f"{organ} apresenta alteração de ecogenicidade.",
                order=idx * 10 + 1
            ),
            TemplateText(
                organ=organ,
                category="finding",
                text=f"{organ} com aumento de dimensões.",
                order=idx * 10 + 2
            )
        ])
    
    # Insert templates
    for template in templates:
        await db.templates.insert_one(template.model_dump())
    
    # Default reference values (simplified example)
    ref_values = [
        # Kidney examples
        ReferenceValue(organ="Rim Esquerdo", measurement_type="comprimento", species="dog", size="small", min_value=3.5, max_value=5.5, unit="cm"),
        ReferenceValue(organ="Rim Esquerdo", measurement_type="comprimento", species="dog", size="medium", min_value=5.0, max_value=7.0, unit="cm"),
        ReferenceValue(organ="Rim Esquerdo", measurement_type="comprimento", species="dog", size="large", min_value=6.5, max_value=9.0, unit="cm"),
        ReferenceValue(organ="Rim Direito", measurement_type="comprimento", species="dog", size="small", min_value=3.5, max_value=5.5, unit="cm"),
        ReferenceValue(organ="Rim Direito", measurement_type="comprimento", species="dog", size="medium", min_value=5.0, max_value=7.0, unit="cm"),
        ReferenceValue(organ="Rim Direito", measurement_type="comprimento", species="dog", size="large", min_value=6.5, max_value=9.0, unit="cm"),
        # Liver
        ReferenceValue(organ="Fígado", measurement_type="espessura", species="dog", size="small", min_value=2.0, max_value=4.0, unit="cm"),
        ReferenceValue(organ="Fígado", measurement_type="espessura", species="dog", size="medium", min_value=3.0, max_value=5.5, unit="cm"),
        ReferenceValue(organ="Fígado", measurement_type="espessura", species="dog", size="large", min_value=4.0, max_value=7.0, unit="cm"),
        # Spleen
        ReferenceValue(organ="Baço", measurement_type="espessura", species="dog", size="small", min_value=0.5, max_value=1.5, unit="cm"),
        ReferenceValue(organ="Baço", measurement_type="espessura", species="dog", size="medium", min_value=1.0, max_value=2.0, unit="cm"),
        ReferenceValue(organ="Baço", measurement_type="espessura", species="dog", size="large", min_value=1.5, max_value=2.5, unit="cm"),
    ]
    
    for ref_value in ref_values:
        await db.reference_values.insert_one(ref_value.model_dump())
    
    return {"message": "Default data initialized successfully"}

# Include the router in the main app
app.include_router(api_router)

app.add_middleware(
    CORSMiddleware,
    allow_credentials=True,
    allow_origins=os.environ.get('CORS_ORIGINS', '*').split(','),
    allow_methods=["*"],
    allow_headers=["*"],
)

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

@app.on_event("shutdown")
async def shutdown_db_client():
    client.close()